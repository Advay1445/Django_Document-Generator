from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

def generate_mediation_docx(data):
    doc = Document()

    # --- 1. PAGE SETUP ---
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Set default font to Times New Roman
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # --- 2. HEADER SECTION ---
    headers = [
        ("FORM 'A'", 11),
        ("MEDIATION APPLICATION FORM", 11),
        ("[REFER RULE 3(1)]", 10),
        ("Mumbai District Legal Services Authority", 11),
        ("City Civil Court, Mumbai", 11)
    ]

    for text, size in headers:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)

    doc.add_paragraph()

    # --- 3. THE MAIN TABLE ---
    # The original has 3 columns: Index, Label, and Data
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    
    # Set exact column widths
    table.columns[0].width = Inches(0.4)
    table.columns[1].width = Inches(2.3)
    table.columns[2].width = Inches(4.3)

    def add_row(col1, col2, col3, is_bold=False):
        row = table.add_row().cells
        row[0].text = str(col1)
        row[1].text = str(col2)
        row[2].text = str(col3)
        if is_bold:
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    def add_full_width_header(text):
        row = table.add_row().cells
        merged = row[0].merge(row[1]).merge(row[2])
        run = merged.paragraphs[0].add_run(text)
        run.bold = True
        merged.paragraphs[0].paragraph_format.space_before = Pt(3)
        merged.paragraphs[0].paragraph_format.space_after = Pt(3)

    # --- CONSTRUCTION ---
    add_full_width_header("DETAILS OF PARTIES:")

    # 1. Applicant Section
    add_row("1", "Name of Applicant", data.get('client_name', ''), is_bold=True)
    
    row_addr_label = table.add_row().cells
    merged_addr = row_addr_label[0].merge(row_addr_label[1]).merge(row_addr_label[2])
    merged_addr.text = "Address and contact details of Applicant"
    merged_addr.paragraphs[0].runs[0].bold = True

    add_row("1", "Address", f"REGISTERED ADDRESS:\n{data.get('branch_address', '')}\n\nCORRESPONDENCE BRANCH ADDRESS:\n{data.get('branch_address_corr', '')}", is_bold=True)
    add_row("", "Telephone No.", data.get('telephone', ''), is_bold=True)
    add_row("", "Mobile No.", data.get('mobile', ''), is_bold=True)
    add_row("", "Email ID", data.get('email', ''), is_bold=True)

    # 2. Opposite Party Section
    add_row("2", "Name, Address and Contact details of Opposite Party:", "", is_bold=True)
    
    row_op_label = table.add_row().cells
    merged_op = row_op_label[0].merge(row_op_label[1]).merge(row_op_label[2])
    merged_op.text = "Address and contact details of Defendant/s"
    merged_op.paragraphs[0].runs[0].bold = True

    add_row("", "Name", data.get('customer_name', ''), is_bold=True)
    add_row("", "Address", f"REGISTERED ADDRESS:\n{data.get('address1', '')}\n\nCORRESPONDENCE ADDRESS:\n{data.get('address_corr', '')}", is_bold=True)
    add_row("", "Telephone No.", data.get('op_telephone', ''), is_bold=True)
    add_row("", "Mobile No.", data.get('op_mobile', ''), is_bold=True)
    add_row("", "Email ID", data.get('op_email', ''), is_bold=True)

    # --- 3. DISPUTE SECTION (SINGLE COMPARTMENTS) ---
    
    # Header Row: "DETAILS OF DISPUTE:"
    header_row = table.add_row().cells
    # Merge all 3 columns into one compartment
    merged_header = header_row[0].merge(header_row[1]).merge(header_row[2])
    merged_header.text = "DETAILS OF DISPUTE:"
    merged_header.paragraphs[0].runs[0].bold = True
    merged_header.paragraphs[0].paragraph_format.space_before = Pt(6)
    merged_header.paragraphs[0].paragraph_format.space_after = Pt(6)

    # Rule Row: Underlined text
    rule_row = table.add_row().cells
    merged_rule = rule_row[0].merge(rule_row[1]).merge(rule_row[2])
    merged_rule.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_rule = merged_rule.paragraphs[0].add_run("THE COMM. COURTS (PRE-INSTITUTION.........SETTLEMENT) RULES,2018")
    run_rule.bold = True
    run_rule.underline = True
    merged_rule.paragraphs[0].paragraph_format.space_before = Pt(6)
    merged_rule.paragraphs[0].paragraph_format.space_after = Pt(6)

    # Final Nature of Dispute Row
    nature_row = table.add_row().cells
    # Merge all 3 columns into one compartment
    merged_nature = nature_row[0].merge(nature_row[1]).merge(nature_row[2])
    nature_text = "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
    merged_nature.text = nature_text
    merged_nature.paragraphs[0].runs[0].bold = True
    merged_nature.paragraphs[0].paragraph_format.space_before = Pt(6)
    merged_nature.paragraphs[0].paragraph_format.space_after = Pt(6)

    # --- 4. EXPORT ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer